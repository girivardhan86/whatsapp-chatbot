from langchain_community.document_loaders import TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import Chroma
from docx import Document as DocxDocument
from langchain_core.documents import Document

import os
import pytesseract
from PIL import Image
import requests
from bs4 import BeautifulSoup

# ================= GLOBALS =================
EMBED_MODEL = "sentence-transformers/all-MiniLM-L6-v2"

# ✅ LOAD EMBEDDING ONLY ONCE
embedding_model = HuggingFaceEmbeddings(
    model_name=EMBED_MODEL
)

# ================= DB =================
def get_db(user_id):

    persist_dir = f"db/{user_id}"
    os.makedirs(persist_dir, exist_ok=True)

    return Chroma(
        persist_directory=persist_dir,
        embedding_function=embedding_model
    )

# ================= SPLITTER =================
def get_splitter():
    return RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50
    )

# ================= FILE INGEST =================
def ingest_file(file_path, user_id="default"):

    docs = []

    if file_path.endswith(".txt"):
        docs = TextLoader(
            file_path,
            encoding="utf-8"
        ).load()

    elif file_path.endswith(".pdf"):
        docs = PyPDFLoader(file_path).load()

    elif file_path.endswith(".docx"):

        doc = DocxDocument(file_path)

        text = "\n".join(
            [p.text for p in doc.paragraphs]
        )

        docs = [
            Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "type": "docx"
                }
            )
        ]

    else:
        print("❌ Unsupported file")
        return

    _store_docs(docs, user_id)

    print(f"✅ File indexed for user: {user_id}")

# ================= IMAGE OCR =================
def ingest_image(file_path, user_id="default"):

    try:
        # Windows path if needed:
        # pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

        img = Image.open(file_path)

        text = pytesseract.image_to_string(img)

        docs = [
            Document(
                page_content=text,
                metadata={
                    "source": file_path,
                    "type": "image"
                }
            )
        ]

        _store_docs(docs, user_id)

        print(f"✅ Image indexed for user: {user_id}")

    except Exception as e:
        print("❌ OCR Error:", e)

# ================= URL INGEST =================
def ingest_url(url, user_id="default"):

    try:
        res = requests.get(url, timeout=10)

        soup = BeautifulSoup(
            res.text,
            "html.parser"
        )

        text = soup.get_text(separator="\n")

        docs = [
            Document(
                page_content=text,
                metadata={
                    "source": url,
                    "type": "url"
                }
            )
        ]

        _store_docs(docs, user_id)

        print(f"✅ URL indexed for user: {user_id}")

    except Exception as e:
        print("❌ URL Error:", e)

# ================= STORE DOCS =================
def _store_docs(docs, user_id):

    splitter = get_splitter()

    split_docs = splitter.split_documents(docs)

    db = get_db(user_id)

    db.add_documents(split_docs)

    print("✅ Stored in vector DB")

# ================= RETRIEVER =================
def get_retriever(user_id="default"):

    db = get_db(user_id)

    return db.as_retriever(
        search_kwargs={"k": 3}
    )